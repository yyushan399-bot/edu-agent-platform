"""DOCX 解析：提取段落与表格，支持分块输出。"""

from __future__ import annotations

from pathlib import Path
from typing import Literal, TypedDict

DocxBlockType = Literal["text", "table"]


def _load_docx(path: Path):
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ImportError(
            "处理 Word 需要安装 python-docx：pip install python-docx"
        ) from exc
    return DocxDocument(path)


class DocxBlock(TypedDict):
    """DOCX 内容块。"""

    type: DocxBlockType
    content: str
    index: int


def _table_to_tsv(table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [(cell.text or "").strip() for cell in row.cells]
        if any(cells):
            rows.append("\t".join(cells))
    return "\n".join(rows)


def parse_docx_blocks(docx_path: str | Path) -> list[DocxBlock]:
    """按顺序提取段落与表格块。"""
    path = Path(docx_path).resolve()
    if not path.is_file():
        raise FileNotFoundError(f"DOCX 不存在: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"仅支持 .docx 文件: {path}")

    doc = _load_docx(path)
    blocks: list[DocxBlock] = []
    index = 0

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            index += 1
            blocks.append({"type": "text", "content": text, "index": index})

    for table in doc.tables:
        tsv = _table_to_tsv(table).strip()
        if tsv:
            index += 1
            blocks.append({"type": "table", "content": tsv, "index": index})

    return blocks


def parse_docx_text(docx_path: str | Path) -> str:
    """从 Word 文档提取纯文本（段落 + 表格）。"""
    blocks = parse_docx_blocks(docx_path)
    if not blocks:
        return ""
    return "\n\n".join(block["content"] for block in blocks)
