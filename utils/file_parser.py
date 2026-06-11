"""从上传文件字节流中提取纯文本。"""

from __future__ import annotations

import io
from pathlib import Path

SUPPORTED_SUFFIXES = frozenset({".pdf", ".docx", ".txt"})


def _suffix_from_filename(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gbk", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore").strip()


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [(cell.text or "").strip() for cell in row.cells if (cell.text or "").strip()]
            if row_text:
                parts.append("\t".join(row_text))

    return "\n".join(parts).strip()


def _extract_pdf(file_bytes: bytes) -> str:
    import fitz

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        pages: list[str] = []
        for page_index in range(len(doc)):
            page = doc[page_index]
            page_text = (page.get_text("text") or "").strip()
            if page_text:
                pages.append(f"\n--- 第 {page_index + 1} 页 ---\n{page_text}")
        return "\n".join(pages).strip()
    finally:
        doc.close()


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    从文件字节内容中提取纯文本。

    Args:
        file_bytes: 上传文件的原始字节。
        filename: 原始文件名，用于判断扩展名。

    Returns:
        提取出的纯文本。

    Raises:
        ValueError: 文件为空、类型不支持或未能提取文本。
    """
    if not file_bytes:
        raise ValueError("文件内容为空。")

    suffix = _suffix_from_filename(filename)
    if suffix not in SUPPORTED_SUFFIXES:
        allowed = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise ValueError(f"暂不支持的文件类型：{suffix or '(无扩展名)'}。当前支持 {allowed}")

    if suffix == ".txt":
        text = _extract_txt(file_bytes)
    elif suffix == ".docx":
        text = _extract_docx(file_bytes)
    elif suffix == ".pdf":
        text = _extract_pdf(file_bytes)
    else:
        raise ValueError(f"暂不支持的文件类型：{suffix}")

    if not text.strip():
        raise ValueError("未能从文件中提取有效文本。")

    return text


__all__ = ["SUPPORTED_SUFFIXES", "extract_text_from_file"]
