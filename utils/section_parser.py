"""整份 PBL 报告 → 7 章节切分（学生反馈 · 阶段 1）。"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from utils.section_constants import SECTION_NAMES

MIN_SECTION_CHARS = 10

# 章节标题别名（整行匹配，不含正文）
SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "文献检索": ("文献", "文献综述", "文献调研", "文献回顾"),
    "问题提出": ("问题", "研究背景", "研究问题", "课题背景"),
    "理论分析": ("理论", "模型", "机理分析", "理论模型"),
    "数值模拟": ("模拟", "编程", "仿真", "数值仿真", "计算机模拟"),
    "实验实施": ("实验", "实验验证", "实验方案", "实验过程"),
    "数据分析": ("数据", "数据处理", "结果分析", "数据结果"),
    "结论生成": ("结论", "创新点", "总结与展望", "总结", "展望"),
}

_CN_NUM_MAP = {
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
}


@dataclass
class SectionChunk:
    section_name: str
    text: str
    char_count: int
    title_line: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "section_name": self.section_name,
            "text": self.text,
            "char_count": self.char_count,
            "title_line": self.title_line,
        }


@dataclass
class SectionParseResult:
    sections: list[SectionChunk] = field(default_factory=list)
    unmatched_text: str = ""
    warnings: list[str] = field(default_factory=list)

    @property
    def found_sections(self) -> list[str]:
        return [s.section_name for s in self.sections]

    @property
    def missing_sections(self) -> list[str]:
        found = set(self.found_sections)
        return [name for name in SECTION_NAMES if name not in found]

    def get_section_text(self, section_name: str) -> str | None:
        for chunk in self.sections:
            if chunk.section_name == section_name:
                return chunk.text
        return None

    def to_dict(self) -> dict[str, Any]:
        return {
            "sections": [s.to_dict() for s in self.sections],
            "unmatched_text": self.unmatched_text,
            "warnings": self.warnings,
            "found_sections": self.found_sections,
            "missing_sections": self.missing_sections,
        }


def _heading_index(line: str) -> int | None:
    stripped = line.strip()
    match = re.match(r"^([一二三四五六七八九十]+|\d+)", stripped)
    if not match:
        return None
    token = match.group(1)
    if token.isdigit():
        return int(token)
    if len(token) == 1 and token in _CN_NUM_MAP:
        return _CN_NUM_MAP[token]
    if token == "十":
        return 10
    return None


def _build_title_pattern(keyword: str) -> re.Pattern[str]:
    return re.compile(
        r"^(?:第?\s*[一二三四五六七八九十\d]+[\s\.、,，\)）]?\s*)?"
        + re.escape(keyword)
        + r"\s*$",
        re.IGNORECASE,
    )


def _build_alias_pattern(alias: str) -> re.Pattern[str]:
    return re.compile(
        r"^(?:第?\s*[一二三四五六七八九十\d]+[\s\.、,，\)）]?\s*)?"
        + re.escape(alias)
        + r"\s*$",
        re.IGNORECASE,
    )


def _resolve_ambiguous_alias(alias: str, heading_idx: int | None) -> str | None:
    if heading_idx is None:
        return None
    if heading_idx <= len(SECTION_NAMES):
        candidate = SECTION_NAMES[heading_idx - 1]
        if alias in SECTION_ALIASES.get(candidate, ()):
            return candidate
    return None


def match_section_title(line: str) -> tuple[str | None, str]:
    """识别单行是否为章节标题，返回 (章节名, 原始标题行)。"""
    para = line.strip()
    if not para:
        return None, ""

    heading_idx = _heading_index(para)

    for section_name in sorted(SECTION_NAMES, key=len, reverse=True):
        if _build_title_pattern(section_name).match(para):
            return section_name, para

    alias_hits: list[tuple[int, str, str]] = []
    for section_name in SECTION_NAMES:
        for alias in SECTION_ALIASES.get(section_name, ()):
            if _build_alias_pattern(alias).match(para):
                alias_hits.append((len(alias), section_name, alias))

    if not alias_hits:
        return None, ""

    alias_hits.sort(key=lambda item: item[0], reverse=True)
    if len(alias_hits) == 1:
        return alias_hits[0][1], para

    resolved = _resolve_ambiguous_alias(alias_hits[0][2], heading_idx)
    if resolved:
        return resolved, para

    for _, section_name, _ in alias_hits:
        if heading_idx is not None and SECTION_NAMES[heading_idx - 1] == section_name:
            return section_name, para

    return alias_hits[0][1], para


def split_report_paragraphs(
    paragraphs: list[str],
    *,
    file_path: str = "",
    min_section_chars: int = MIN_SECTION_CHARS,
) -> SectionParseResult:
    """按段落列表切分 7 章节（DOCX 推荐）。"""
    warnings: list[str] = []
    preamble: list[str] = []
    sections_raw: list[dict[str, Any]] = []

    current_section: str | None = None
    current_title = ""
    current_content: list[str] = []

    for para in paragraphs:
        text = para.strip()
        if not text:
            continue

        matched_section, title_line = match_section_title(text)
        if matched_section:
            if current_section and current_content:
                content = "\n".join(current_content).strip()
                if len(content) >= min_section_chars:
                    sections_raw.append(
                        {
                            "section": current_section,
                            "content": content,
                            "title_line": current_title,
                        }
                    )
                else:
                    warnings.append(
                        f"章节「{current_section}」内容过短（{len(content)} 字），已忽略。"
                    )
            elif current_section is None and current_content:
                preamble.extend(current_content)

            current_section = matched_section
            current_title = title_line
            current_content = []
        elif current_section:
            current_content.append(text)
        else:
            preamble.append(text)

    if current_section and current_content:
        content = "\n".join(current_content).strip()
        sections_raw.append(
            {
                "section": current_section,
                "content": content,
                "title_line": current_title,
            }
        )

    result = SectionParseResult(
        unmatched_text="\n".join(preamble).strip(),
        warnings=warnings,
    )
    result.sections = _finalize_sections(sections_raw, paragraphs, warnings, min_section_chars)
    return result


def split_report_text(
    text: str,
    *,
    file_path: str = "",
    min_section_chars: int = MIN_SECTION_CHARS,
) -> SectionParseResult:
    """按行切分整份报告文本（PDF / TXT）。"""
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    return split_report_paragraphs(lines, file_path=file_path, min_section_chars=min_section_chars)


def _finalize_sections(
    sections_raw: list[dict[str, Any]],
    paragraphs: list[str],
    warnings: list[str],
    min_section_chars: int,
) -> list[SectionChunk]:
    section_best: dict[str, dict[str, Any]] = {}
    for item in sections_raw:
        name = item["section"]
        if name not in section_best or len(item["content"]) > len(section_best[name]["content"]):
            section_best[name] = item

    found = set(section_best)
    missing = [name for name in SECTION_NAMES if name not in found]
    if missing:
        warnings.append(f"未找到章节：{', '.join(missing)}")
        full_text = "\n".join(p.strip() for p in paragraphs if p.strip())
        for sec in missing:
            idx = full_text.find(sec)
            if idx == -1:
                continue
            next_idx = len(full_text)
            for other in SECTION_NAMES:
                if other == sec:
                    continue
                other_idx = full_text.find(other, idx + len(sec))
                if other_idx != -1 and other_idx < next_idx:
                    next_idx = other_idx
            content = full_text[idx + len(sec) : next_idx].strip()
            if len(content) >= min_section_chars:
                section_best[sec] = {
                    "section": sec,
                    "content": content,
                    "title_line": sec,
                }
                warnings.append(f"回退查找命中章节：{sec}（{len(content)} 字）")

    order_map = {name: index for index, name in enumerate(SECTION_NAMES)}
    ordered = sorted(section_best.values(), key=lambda item: order_map.get(item["section"], 999))

    chunks: list[SectionChunk] = []
    for item in ordered:
        content = item["content"]
        if len(content) < min_section_chars:
            continue
        chunks.append(
            SectionChunk(
                section_name=item["section"],
                text=content,
                char_count=len(content),
                title_line=item.get("title_line", ""),
            )
        )
    return chunks


def read_docx_paragraphs(path: str | Path) -> list[str]:
    """读取 DOCX 段落（保留标题行结构，不含表格）。"""
    from docx import Document

    doc = Document(str(path))
    return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]


def read_docx_paragraphs_from_bytes(file_bytes: bytes) -> list[str]:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    return [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]


def parse_report_from_bytes(file_bytes: bytes, filename: str) -> SectionParseResult:
    """从上传文件字节解析并切分章节。"""
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".docx":
        paragraphs = read_docx_paragraphs_from_bytes(file_bytes)
        return split_report_paragraphs(paragraphs, file_path=filename)
    from utils.file_parser import extract_text_from_file

    text = extract_text_from_file(file_bytes, filename)
    return split_report_text(text, file_path=filename)


def parse_report_from_path(path: str | Path) -> SectionParseResult:
    """从本地文件路径解析并切分章节。"""
    file_path = Path(path)
    if not file_path.is_file():
        raise FileNotFoundError(f"报告文件不存在：{file_path}")

    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return split_report_paragraphs(
            read_docx_paragraphs(file_path),
            file_path=str(file_path),
        )

    from utils.file_parser import extract_text_from_path

    text = extract_text_from_path(file_path)
    return split_report_text(text, file_path=str(file_path))


__all__ = [
    "MIN_SECTION_CHARS",
    "SectionChunk",
    "SectionParseResult",
    "match_section_title",
    "parse_report_from_bytes",
    "parse_report_from_path",
    "read_docx_paragraphs",
    "read_docx_paragraphs_from_bytes",
    "split_report_paragraphs",
    "split_report_text",
]
